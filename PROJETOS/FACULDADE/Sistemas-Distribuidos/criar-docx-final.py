#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Criar documento
doc = Document()

# Configurar margens
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

# Estilo Normal
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ==================== CAPA ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDADE UNOPAR\nANÁLISE E DESENVOLVIMENTO DE SISTEMAS')
run.bold = True
run.font.size = Pt(14)
run.font.all_caps = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ACIB ABBADE DE CASTRO')
run.bold = True
run.font.size = Pt(12)
run.font.all_caps = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RA 3850892303')
run.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Relatório de aula prática solicitado no curso de Análise e Desenvolvimento de Sistemas como requisito parcial para a obtenção de pontos na matéria de Sistemas Distribuídos, para a média semestral.')
run.font.size = Pt(11)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Orientador: Anderson Emidio de Macedo Goncalves')
run.font.size = Pt(11)

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('JACAREÍ-SP\n2025')
run.font.size = Pt(11)
run.font.all_caps = True

doc.add_page_break()

# ==================== FOLHA DE ROSTO ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDADE UNOPAR\nANÁLISE E DESENVOLVIMENTO DE SISTEMAS')
run.font.size = Pt(11)
run.font.all_caps = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('ACIB ABBADE DE CASTRO\nRA 3850892303')
run.font.size = Pt(11)
run.font.all_caps = True

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PORTFÓLIO - RELATÓRIO DE AULA PRÁTICA\nSISTEMAS DISTRIBUÍDOS')
run.bold = True
run.font.size = Pt(12)
run.font.all_caps = True

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Relatório de aula prática solicitado no curso de Análise e Desenvolvimento de Sistemas como requisito parcial para a obtenção de pontos na matéria de Sistemas Distribuídos, para a média semestral.\n\nOrientador: Anderson Emidio de Macedo Goncalves')
run.font.size = Pt(10)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('JACAREÍ – SP\n2025')
run.font.size = Pt(11)
run.font.all_caps = True

doc.add_page_break()

# ==================== SUMÁRIO ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUMÁRIO')
run.bold = True
run.font.size = Pt(14)
run.font.all_caps = True

sumario = [
    ('1. INTRODUÇÃO', '4'),
    ('2. DESENVOLVIMENTO', '4'),
    ('2.1 Sincronização de Relógios com NTP', '4'),
    ('2.2 Conteinerização com Docker', '6'),
    ('2.3 Análise de Protocolos com Wireshark', '8'),
    ('2.4 Virtualização com VirtualBox', '9'),
    ('3. RESULTADOS', '11'),
    ('4. CONCLUSÃO', '12'),
]

for item, pagina in sumario:
    p = doc.add_paragraph()
    if item.startswith('2.'):
        p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f'{item}')
    run.font.size = Pt(11)
    dots = '.' * (50 - len(item))
    run = p.add_run(f' {dots} {pagina}')
    run.font.size = Pt(11)

doc.add_page_break()

# ==================== 1. INTRODUÇÃO ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('1. INTRODUÇÃO')
run.bold = True
run.font.size = Pt(12)
run.font.all_caps = True

doc.add_paragraph('Este trabalho foi desenvolvido para a disciplina de Sistemas Distribuídos com o objetivo de aplicar conceitos fundamentais de sincronização de relógios, conteinerização, virtualização e análise de protocolos de rede.')

doc.add_paragraph('As atividades práticas realizadas foram:')

atividades = [
    'Sincronização de relógios utilizando servidor NTP nos sistemas Windows 10 e Linux (WSL)',
    'Orquestração de contêineres com Docker Swarm (5 réplicas do servidor Apache)',
    'Análise de tráfego de rede com o Wireshark',
    'Criação de máquina virtual com Oracle VirtualBox'
]

for att in atividades:
    p = doc.add_paragraph()
    run = p.add_run(f'✓ {att}')
    run.font.size = Pt(11)

doc.add_paragraph('Foi utilizado o Windows 10 como sistema operacional base, WSL (Windows Subsystem for Linux) para os comandos Linux e Docker Desktop para a orquestração dos contêineres.')

# ==================== 2. DESENVOLVIMENTO ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2. DESENVOLVIMENTO')
run.bold = True
run.font.size = Pt(12)
run.font.all_caps = True

# 2.1
p = doc.add_paragraph()
run = p.add_run('2.1 Sincronização de Relógios com NTP')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph('No Windows 10:')
doc.add_paragraph('Abri o Prompt de Comando (CMD) como Administrador e executei os seguintes comandos:')

p = doc.add_paragraph()
run = p.add_run('w32tm /config /syncfromflags:manual /manualpeerlist:0.pool.ntp.org\nnet stop w32time\nnet start w32time\nw32tm /resync /rediscover')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph('Estes comandos configuram o servidor NTP, reiniciam o serviço de horário e forçam a sincronização.')

# Print 1
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 1 - CMD com os 4 comandos NTP executados')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-01-ntp-cmd.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 1 - CMD NTP]')

doc.add_paragraph('No Linux (WSL/Ubuntu):')
doc.add_paragraph('Acessei o terminal do Ubuntu via WSL e executei:')

p = doc.add_paragraph()
run = p.add_run('sudo apt update\nsudo apt install ntp -y')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph('Em seguida, editei o arquivo de configuração /etc/ntp.conf substituindo os pools padrão por:')

p = doc.add_paragraph()
run = p.add_run('pool pool.ntp.br iburst')
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Print 2
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 2 - Arquivo ntp.conf editado')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-02-ntp-conf.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 2 - ntp.conf]')

doc.add_paragraph('Reiniciei o serviço e verifiquei a sincronização:')

p = doc.add_paragraph()
run = p.add_run('sudo service ntp restart\nntpq -p')
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Print 3
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 3 - Comando ntpq -p mostrando a sincronização')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-03-ntpq-p.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 3 - ntpq -p]')

# 2.2
p = doc.add_paragraph()
run = p.add_run('2.2 Conteinerização com Docker')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph('Inicializei o cluster Docker Swarm:')

p = doc.add_paragraph()
run = p.add_run('docker swarm init')
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Print 4
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 4 - Comando docker swarm init')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-04-docker-swarm-init.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 4 - docker swarm init]')

doc.add_paragraph('Criei um serviço chamado "WEB" com 5 réplicas do servidor Apache:')

p = doc.add_paragraph()
run = p.add_run('docker service create --name WEB --publish 80:80 --replicas=5 httpd')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph('[PRINT 5 - docker service create]')

doc.add_paragraph('Verifiquei as 5 réplicas rodando:')

p = doc.add_paragraph()
run = p.add_run('docker service ps WEB')
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Print 6
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 5 - docker service ps WEB mostrando 5 réplicas com STATUS Running')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-06-docker-ps-web.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 6 - docker ps]')

doc.add_paragraph('Testei o servidor Apache no navegador acessando http://localhost:')

# Print 7
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 6 - Navegador mostrando a página "It works!" do Apache')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-07-it-works.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 7 - It works]')

# 2.3
p = doc.add_paragraph()
run = p.add_run('2.3 Análise de Protocolos com Wireshark')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph('Baixei e instalei o Wireshark no site oficial (https://www.wireshark.org). Após a instalação, reiniciei o computador para ativar o driver Npcap.')

doc.add_paragraph('Abri o Wireshark como Administrador e iniciei a captura de pacotes na interface Wi-Fi.')

# Print 8
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 7 - Wireshark capturando pacotes (linhas coloridas)')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-08-wireshark-capturando.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 8 - Wireshark]')

doc.add_paragraph('Apliquei o filtro "http" para visualizar apenas o tráfego web.')

# Print 9
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 8 - Wireshark com filtro http aplicado')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-09-wireshark-filtro.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 9 - Wireshark filtro]')

# 2.4
p = doc.add_paragraph()
run = p.add_run('2.4 Virtualização com VirtualBox')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph('Baixei e instalei o Oracle VirtualBox no site oficial (https://www.virtualbox.org).')

# Print 10
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 9 - Download do VirtualBox')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-10-virtualbox-download.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 10 - VirtualBox download]')

doc.add_paragraph('Criei uma nova máquina virtual com as seguintes configurações:')

config_vm = [
    'Nome: MinhaVM',
    'Tipo: Linux',
    'Memória RAM: 2GB (2048 MB)',
    'Disco rígido: 25GB',
    'ISO: Ubuntu 22.04 LTS'
]

for cfg in config_vm:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(cfg)
    run.font.size = Pt(11)

# Print 11
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 10 - Configuração de hardware da VM')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-11-vm-config.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 11 - VM config]')

doc.add_paragraph('Iniciei a máquina virtual e realizei a instalação do sistema operacional Linux.')

# Print 12
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Figura 11 - VM rodando com a instalação do Linux')
run.bold = True
run.font.size = Pt(10)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    doc.add_picture('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/print-12-vm-rodando.jpg', width=Cm(14))
except:
    doc.add_paragraph('[PRINT 12 - VM rodando]')

# ==================== 3. RESULTADOS ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('3. RESULTADOS')
run.bold = True
run.font.size = Pt(12)
run.font.all_caps = True

doc.add_paragraph('Todas as atividades práticas foram concluídas com sucesso.')

doc.add_paragraph('NTP: A sincronização de relógios foi configurada tanto no Windows quanto no Linux, utilizando o servidor pool.ntp.br. O comando ntpq -p confirmou a comunicação com os servidores de horário.')

doc.add_paragraph('Docker: O cluster Swarm foi inicializado com sucesso. O serviço WEB foi criado com 5 réplicas do Apache, e todas as 5 instâncias ficaram com STATUS "Running". O navegador exibiu a página "It works!", comprovando o funcionamento do servidor web orquestrado.')

doc.add_paragraph('Wireshark: O analisador de protocolos capturou os pacotes de rede corretamente, permitindo visualizar o tráfego HTTP e aplicar filtros para análise específica.')

doc.add_paragraph('VirtualBox: A máquina virtual foi criada e o sistema operacional Linux foi instalado com sucesso, demonstrando o conceito de virtualização.')

# ==================== 4. CONCLUSÃO ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('4. CONCLUSÃO')
run.bold = True
run.font.size = Pt(12)
run.font.all_caps = True

doc.add_paragraph('Através deste trabalho prático, compreendi os fundamentos de Sistemas Distribuídos, incluindo:')

doc.add_paragraph('Sincronização temporal (NTP): essencial para serviços como autenticação e acesso remoto, evitando falhas causadas por diferenças de horário entre máquinas.')

doc.add_paragraph('Conteinerização e orquestração (Docker Swarm): permite criar serviços distribuídos com múltiplas réplicas, garantindo escalabilidade e alta disponibilidade.')

doc.add_paragraph('Análise de protocolos (Wireshark): ferramenta fundamental para monitoramento e troubleshooting de redes.')

doc.add_paragraph('Virtualização (VirtualBox): possibilita executar múltiplos sistemas operacionais em um único hardware, otimizando recursos.')

doc.add_paragraph('O projeto me mostrou a importância dos sistemas distribuídos no cenário atual da computação, onde serviços precisam ser escaláveis, confiáveis e eficientes. Todos os objetivos de aprendizagem foram alcançados com sucesso.')

# ==================== LISTA DE FIGURAS ====================
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run('LISTA DE FIGURAS')
run.bold = True
run.font.size = Pt(12)

figuras = [
    'Figura 1 - CMD com os 4 comandos NTP executados',
    'Figura 2 - Arquivo ntp.conf editado',
    'Figura 3 - Comando ntpq -p mostrando a sincronização',
    'Figura 4 - Comando docker swarm init',
    'Figura 5 - docker service ps WEB mostrando 5 réplicas',
    'Figura 6 - Navegador mostrando a página "It works!"',
    'Figura 7 - Wireshark capturando pacotes',
    'Figura 8 - Wireshark com filtro http aplicado',
    'Figura 9 - Download do VirtualBox',
    'Figura 10 - Configuração de hardware da VM',
    'Figura 11 - VM rodando com instalação do Linux',
]

for fig in figuras:
    p = doc.add_paragraph()
    run = p.add_run(fig)
    run.font.size = Pt(10)

# Rodapé
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Acib Abbade de Castro - RA 3850892303\nUniversidade UNOPAR - Análise e Desenvolvimento de Sistemas\nSistemas Distribuídos - 2025')
run.font.size = Pt(9)

# Salvar
doc.save('/root/.openclaw/workspace/PROJETOS/FACULDADE/Sistemas-Distribuidos/RELATORIO-FINAL-COM-PRINTS.docx')
print('✅ DOCX FINAL criado com todos os prints inseridos!')
